from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path(__file__).with_name("RunningHub工作台项目需求说明书.docx")
doc = Document(path)
text = "\n".join(p.text for p in doc.paragraphs)
for required in ("基础工作台", "多工作流适配", "软件授权与加密",
                 "有无服务器", "包月会员", "基础版本交付与验收"):
    assert required in text, required
assert len(doc.paragraphs) <= 60
assert len(doc.tables) == 6
assert sum(1 for p in doc.paragraphs if any(
    br.tag == qn("w:br") and br.get(qn("w:type")) == "page"
    for br in p._p.iter(qn("w:br")))) == 1
for table in doc.tables:
    widths = [int(c.get(qn("w:w"))) for c in table._tbl.tblGrid]
    assert sum(widths) == 9504, widths
    assert table._tbl.tblPr.find(qn("w:tblW")).get(qn("w:w")) == "9504"
print({"paragraphs": len(doc.paragraphs), "tables": len(doc.tables),
       "manual_page_breaks": 1, "bytes": path.stat().st_size})
