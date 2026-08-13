from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("RunningHub工作台项目需求说明书.docx")
NAVY, BLUE, GREEN = "173A5E", "2E74B5", "237A57"
INK, MUTED, LINE = "20252B", "66727E", "D8DEE5"
LIGHT, PALE_BLUE, PALE_GREEN = "F2F4F7", "E8EEF5", "EAF5EF"


def rgb(value):
    return RGBColor.from_string(value)


def set_font(run, size=9.5, bold=False, color=INK):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def cell_margins(cell, value=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, amount in (("top", value), ("bottom", value),
                         ("start", 100), ("end", 100)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(amount))
        node.set(qn("w:type"), "dxa")


def borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    group = tbl_pr.find(qn("w:tblBorders"))
    if group is None:
        group = OxmlElement("w:tblBorders")
        tbl_pr.append(group)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), color)
        group.append(node)


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9504")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def para(doc, text, size=9.5, bold=False, color=INK, before=0, after=3,
         align=WD_ALIGN_PARAGRAPH.LEFT, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_with_next = keep
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text):
    return para(doc, text, size=12, bold=True, color=BLUE,
                before=7, after=3, keep=True)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.02
    set_font(p.add_run(text), size=9.2)
    return p


def table(doc, headers, rows, widths, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        shade(t.rows[0].cells[i], PALE_BLUE)
        p = t.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=8.7, bold=True, color=NAVY)
    for data in rows:
        row = t.add_row()
        for i, value in enumerate(data):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(value), size=font_size)
    geometry(t, widths)
    borders(t)
    para(doc, "", after=0)
    return t


def callout(doc, text):
    t = doc.add_table(rows=1, cols=1)
    geometry(t, [9504])
    cell = t.cell(0, 0)
    shade(cell, PALE_GREEN)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    set_font(p.add_run(text), size=9.2, bold=True, color=GREEN)
    para(doc, "", after=0)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.55)
    sec.left_margin = sec.right_margin = Inches(0.7)
    sec.header_distance = sec.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for name in ("List Bullet",):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9.2)

    hp = sec.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("RunningHub 多工作流智能生成工作台"), size=7.5,
             bold=True, color=MUTED)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_font(fp.add_run("客户需求确认稿｜2026-08-13"), size=7.5, color=MUTED)

    para(doc, "RunningHub 多工作流智能生成工作台",
         size=21, bold=True, color=NAVY, after=2, keep=True)
    para(doc, "项目需求说明（简版）", size=11.5, color=MUTED, after=7, keep=True)
    callout(doc, "核心目标：把 RunningHub 上复杂的 ComfyUI 工作流封装成易用工作台。用户只需选择功能、上传素材和提交任务，无需看到节点图或手动操作工作流。基础工作台完整可用是第一优先级。")

    heading(doc, "一、基础工作台")
    table(doc, ["模块", "主要功能"], [
        ["工作流中心", "展示多个图片/视频工作流，支持分类、搜索和选择；未适配完成的功能不可提交。"],
        ["素材上传", "从电脑任意目录选择图片或视频，自动校验必填项并上传到工作流指定节点。"],
        ["账号管理", "添加多个 RunningHub 手机号，显示登录、过期、忙碌等状态；未完成登录的账号可以删除。"],
        ["任务执行", "自动选择空闲账号或使用指定账号；忙碌时排队，完成后自动执行下一任务。"],
        ["任务中心", "显示工作流名称、账号、创建时间、状态、耗时、错误原因和结果文件。"],
        ["结果保存", "按工作流配置保存图片或视频，只保存指定输出节点，避免下载错误结果。"],
    ], [1900, 7604], font_size=8.8)

    heading(doc, "二、多工作流适配")
    para(doc, "上传、提交、等待、完成判断和保存结果做成统一函数；每个工作流只配置名称、工作流 ID、输入节点、输出节点、文件类型和保存动作。新增工作流后使用真实素材完整验证一次。")
    bullet(doc, "当前人物替换工作流：人物图节点 108、动作视频节点 112、输出只保留节点 119。")
    bullet(doc, "保存动作支持 Save Preview、Save Image、Save Video，由不同工作流分别配置。")
    bullet(doc, "工作流发布权限仍由 RunningHub 控制，执行账号必须拥有合法访问权限。")

    heading(doc, "三、RunningHub 登录")
    para(doc, "软件使用自己的登录弹窗。发送短信前需要完成 RunningHub 官方滑块，工作台只同步滑块验证区域，由用户手动拖动，不破解或绕过验证码。滑块通过后立即关闭官方画面，只保留本软件的短信验证码输入框，避免用户看到后面的原始网页。登录状态保存在本机。")

    heading(doc, "四、软件授权与加密")
    bullet(doc, "使用数字签名许可证：发码端保存私钥，客户端只放公钥，用户不能自行伪造授权。")
    bullet(doc, "许可证绑定机器码，并限制允许设备数量，防止同一密钥直接复制给多人。")
    bullet(doc, "登录状态和授权缓存使用系统级加密保存；程序可增加打包、混淆和完整性校验。")
    bullet(doc, "正常用户看不到工作流 JSON 和节点图，但任何本地软件都不能承诺绝对无法破解。")

    doc.add_page_break()
    heading(doc, "五、有无服务器的激活方式区别")
    table(doc, ["比较项", "无服务器授权", "有服务器会员"], [
        ["适合场景", "客户少、永久授权、接受人工处理", "月卡/季卡、续费、设备管理和长期运营"],
        ["激活方式", "客户先提供机器码，再生成该设备专用许可证", "用户输入卡密，服务器首次激活时绑定设备"],
        ["包月使用", "许可证写到期日；续费需要重新发码", "后台延长到期时间，用户通常无需重新激活"],
        ["设备限制", "生成时固定设备，换机需重新发码", "可限制设备数，并在后台解绑或换绑"],
        ["封禁与撤销", "发出后难以远程撤销", "可实时禁用、恢复或封禁"],
        ["时间安全", "可能受到修改系统时间影响", "以服务器时间为准，更可靠"],
        ["离线使用", "完全离线可用", "可设置 1-7 天离线宽限期"],
        ["成本", "开发和维护成本较低", "需要服务器、域名、HTTPS、数据库和维护"],
    ], [1650, 3877, 3977], font_size=8.2)
    callout(doc, "建议：只销售少量永久授权可选无服务器方案；需要包月、续费、封禁、换机和统一管理时，应直接选择服务器方案。纯离线方案无法可靠实现严格包月和实时控制。")

    heading(doc, "六、包月会员使用方式（服务器方案）")
    bullet(doc, "客户直接向运营方付款，运营方在后台创建月卡、季卡、年卡或永久卡并发送卡密，不包含在线支付。")
    bullet(doc, "用户首次输入卡密后绑定设备；软件启动时及每 12-24 小时验证一次，建议设置 3 天离线宽限。")
    bullet(doc, "续费时后台延长原授权；到期后禁止新建任务，但不删除本地历史和已经生成的结果。")
    bullet(doc, "后台基础功能：创建卡密、查看状态、续期、禁用、恢复、设备解绑及操作记录。")
    para(doc, "说明：软件会员只控制本工作台的使用权，不包含 RunningHub 会员、RH 币、算力或其他平台费用。", size=8.8, bold=True, color=NAVY, after=4)

    heading(doc, "七、基础版本交付与验收")
    bullet(doc, "工作台可正常启动，账号登录、任意目录上传、工作流选择、任务排队、状态查看和结果保存链路完整。")
    bullet(doc, "每个适配工作流使用真实素材验证输入节点、唯一输出节点、保存方式和结果可用性。")
    bullet(doc, "登录滑块拖动流畅；验证后不显示 RunningHub 原网页，本软件验证码框保持打开。")
    bullet(doc, "授权功能按照最终选定的离线或服务器方案验收；已有本地结果不因到期或删除账号而丢失。")

    heading(doc, "八、客户需确认")
    table(doc, ["确认项", "需要提供/决定"], [
        ["工作流", "首批工作流清单、优先级、展示名称、有效测试素材及唯一输出节点。"],
        ["授权", "无服务器永久授权，或带服务器的月卡/季卡/永久卡。"],
        ["会员规则", "允许设备数、离线宽限、换机次数、到期处理和是否提供试用。"],
        ["平台权限", "确认执行账号能够访问已发布或已授权的 RunningHub 工作流。"],
    ], [1900, 7604], font_size=8.6)
    callout(doc, "实施顺序：先保证基础工作台稳定完整，再批量适配工作流，最后接入选定的授权方案。在线支付、复杂代理体系、云素材库等不纳入基础版本。")

    props = doc.core_properties
    props.title = "RunningHub 多工作流智能生成工作台项目需求说明（简版）"
    props.subject = "工作台、多工作流、授权与包月会员"
    props.author = "项目组"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
